#!/usr/bin/env python3
"""Apply the v0.3.0a2 documentation corrections to the research-note DOCX.

The edit is intentionally narrow and idempotent. It refuses to alter a document
whose expected source paragraphs have drifted, so a future revision cannot be
silently overwritten by this release-maintenance helper.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


REPLACEMENTS: dict[int, tuple[str, str]] = {
    3: (
        "20 July 2026",
        "20 July 2026 - revised for audit engine 0.3.0a2",
    ),
    25: (
        "These structures have been implemented as version 0.2 of a standard-library, exact-rational audit engine. The implementation checks claim manifests, product-valued hard gates, certificate complexes, context transport, path holonomy, observation descent, atomic-concentration records, arithmetic no-go conditions, and exact propagation of quantitative defect bounds. Seventeen known-answer tests pass. This establishes an executable structural kernel. It does not prove the Riemann hypothesis, construct the required infinite-dimensional operator, or validate every scientific claim in the source catalogue.",
        "These structures have been implemented as version 0.3.0a2 of a standard-library, exact-rational audit engine. The implementation checks claim manifests, evidence-derived product-valued hard gates, certificate complexes, context transport, path holonomy, observation descent, atomic-concentration records, arithmetic no-go conditions, schema/runtime parity, null discrimination, and exact propagation of quantitative defect bounds. More than sixty deterministic tests pass. This establishes an executable structural kernel. It does not prove the Riemann hypothesis, construct the required infinite-dimensional operator, or validate every scientific claim in the source catalogue.",
    ),
    252: (
        "The producer does not assign its own theorem verdict. The verifier emits exactly one of PROVEN, REFUTED, SUPPORTED, UNRESOLVED, or MALFORMED, with a witness or missing-obligation list whenever possible.",
        "The producer does not assign its own theorem verdict. The current verifier emits a mechanical audit decision, structured findings, and an actual-execution ledger. A theorem is mechanically admissible only when a claim-bound, hash-verified proof artifact reports a passing result and every hard gate is derived as passing from the complete bound evidence set. Human research classifications such as proven, refuted, supported, or unresolved remain separate from the engine's finite checks unless a declared verification adapter supplies the corresponding certificate.",
    ),
    255: (
        "Every command emits canonical JSON findings. Exit code zero means no blocking finding; exit code one means an error, block, or demotion occurred.",
        "Every command emits canonical JSON findings and an actual-execution ledger. Exit code zero means that no blocking finding was produced; exit code one means that a valid input was blocked or demoted; exit code two means malformed input or command usage; and exit code seventy reports an internal verifier failure. Checks skipped after a fail-closed structural error are recorded as not run, never inferred as executed.",
    ),
    347: (
        "{\n  \"claim\": {\n    \"id\": \"example:claim\",\n    \"type\": \"theorem_schema\",\n    \"epistemic_status\": \"checked\",\n    \"deployment_status\": \"sandboxed\"\n  },\n  \"admission\": {\n    \"hard_gates\": [\"naturality\", \"boundary\", \"atomic_rigidity\"],\n    \"gate_results\": [\n      {\"id\": \"naturality\", \"state\": \"pass\", \"fatal\": true},\n      {\"id\": \"boundary\", \"state\": \"unrun\", \"fatal\": true},\n      {\"id\": \"atomic_rigidity\", \"state\": \"pass\", \"fatal\": true}\n    ]\n  },\n  \"dependency_graph\": {\n    \"root\": \"example:claim\",\n    \"nodes\": [\"naturality\", \"boundary\", \"atomic_rigidity\", \"example:claim\"],\n    \"edges\": [\n      {\"source\": \"naturality\", \"target\": \"example:claim\"},\n      {\"source\": \"boundary\", \"target\": \"example:claim\"},\n      {\"source\": \"atomic_rigidity\", \"target\": \"example:claim\"}\n    ]\n  }\n}",
        "{\n  \"manifest_version\": \"0.3.0\",\n  \"draft\": true,\n  \"claim\": {\n    \"id\": \"example:claim\",\n    \"title\": \"Example research claim\",\n    \"type\": \"conjecture\",\n    \"evidence_maturity\": \"declared\",\n    \"deployment_status\": \"research_only\",\n    \"statement\": \"A precisely scoped example claim remains unresolved.\",\n    \"scope\": \"illustrative manifest only\"\n  },\n  \"system\": {\"domain\": \"declared system\", \"state_type\": \"typed states\"},\n  \"observation\": {\"kernel_or_instrument\": \"declared instrument\", \"legal_filtration\": {}},\n  \"representation\": {\"kind\": \"identity\"},\n  \"target\": {\"outcome\": \"declared outcome\", \"horizon\": \"declared horizon\", \"loss_or_score\": \"declared score\"},\n  \"experiment\": {\"baseline_model\": \"ordinary baseline\", \"search_budget\": \"not started\"},\n  \"admission\": {\n    \"hard_gates\": [\"naturality\"],\n    \"gate_results\": [{\"id\": \"naturality\", \"state\": \"unrun\", \"fatal\": true, \"evidence\": []}]\n  },\n  \"demotion\": {\n    \"owner\": \"maintainer\",\n    \"rules\": [{\"if\": \"counterexample\", \"then\": \"retire\"}],\n    \"negative_result_destination\": \"negative-results/\"\n  },\n  \"preservation\": {\"known_failures\": [\"No proof or experiment has been supplied.\"]},\n  \"evidence\": []\n}",
    ),
    349: (
        "The local version 0.2 implementation was tested under Python 3.12 with seventeen deterministic unit tests. The suite covers exact rational parsing, rejection of floats, matrix rank and products, chain-complex validation, natural and broken transports, exact residual witnesses, observation descent, query-family kernels, manifest demotion rules, finite prime-comb rejection, admission with unrun gates, fatal dependency propagation, atomic-modulus consistency and evasion, affine-defect associativity, and understatement rejection.",
        "The version 0.3.0a2 engine is tested under pinned Python 3.12.13 and setuptools 82.0.1. Its 62 deterministic unit tests cover exact arithmetic, schema/runtime parity, evidence-derived gate states, conflict and omission detection, failed-proof rejection, fail-closed arithmetic configuration, fatal propagation, and certificate calculations. A dedicated Null-Discrimination suite preserves the four known false-pass cases as regression fixtures.",
    ),
    350: (
        "This record is a build fact, not an independent replication. A public continuous-integration run and a second verifier remain required for external-replication status.",
        "Continuous integration is a build fact, not independent replication; the latter requires an independently operated verifier and claim-bound passing evidence.",
    ),
}


ALTERNATE_SOURCES: dict[int, tuple[str, ...]] = {
    349: (
        "The version 0.3.0a2 implementation is tested under pinned Python 3.12.13 and setuptools 82.0.1 with more than sixty deterministic unit tests. The suite covers exact rational parsing, rejection of floats, matrix rank and products, chain-complex validation, natural and broken transports, exact residual witnesses, observation descent, query-family kernels, schema/runtime parity, evidence-derived gate states, conflict and omission detection, failed-proof rejection, arithmetic-plugin fail-closed behavior, finite prime-comb rejection, fatal dependency propagation, atomic-modulus consistency and evasion, affine-defect associativity, and understatement rejection. A dedicated Null-Discrimination suite preserves the known false-pass cases as regression fixtures.",
    ),
    350: (
        "Local and public continuous-integration runs are build facts, not independent scientific replication. External-replication status still requires an independently operated verifier, independently obtained artifacts where the claim demands them, and an explicit passing replication record bound to the claim.",
    ),
}


TABLE_REPLACEMENTS: tuple[tuple[int, int, int, str, str], ...] = (
    (0, 8, 1, "Plausible but unresolved computationally", "Plausible; unresolved"),
    (0, 9, 1, "Plausible but unresolved computationally", "Plausible; unresolved"),
)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "path",
        nargs="?",
        type=Path,
        default=Path("research/Audit_Descent_Calculus.docx"),
    )
    args = parser.parse_args()

    document = Document(args.path)
    for index, (old, new) in REPLACEMENTS.items():
        paragraph = document.paragraphs[index]
        if paragraph.text == new:
            continue
        accepted_sources = (old, *ALTERNATE_SOURCES.get(index, ()))
        if paragraph.text not in accepted_sources:
            raise SystemExit(
                f"paragraph {index} has drifted; refusing replacement: {paragraph.text!r}"
            )
        paragraph.text = new

    for table_index, row_index, cell_index, old, new in TABLE_REPLACEMENTS:
        cell = document.tables[table_index].cell(row_index, cell_index)
        if cell.text == new:
            continue
        if cell.text != old:
            raise SystemExit(
                "table cell has drifted; refusing replacement: "
                f"table={table_index} row={row_index} cell={cell_index} text={cell.text!r}"
            )
        cell.text = new

    document.save(args.path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
